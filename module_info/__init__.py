import azure.functions as func
import logging
from langchain.document_loaders import PyPDFLoader, DirectoryLoader, TextLoader, JSONLoader
from langchain_community.embeddings import AzureOpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
import faiss
from langchain.schema import Document
from openai import OpenAI
from docx import Document as DocxDocument
import os
import io
import json
import tempfile
import pickle
from azure.storage.blob import BlobServiceClient
from openai import AzureOpenAI
from datetime import datetime

os.environ['FAISS_NO_GPU'] = '1'

logger = logging.getLogger(__name__)

class KnowledgeBase:
    def __init__(self, text_data):
        self.embeddings = AzureOpenAIEmbeddings(
            azure_deployment="text-embedding-ada-002",  # Nom de votre déploiement d'embeddings
            openai_api_key=os.environ["AZURE_OPENAI_API_KEY"],
            azure_endpoint=os.environ["AZURE_OPENAI_ENDPOINT"],
            openai_api_version="2023-05-15"  ,chunk_size=1
        )
        self.text_splitter = None 
        self.documents=self.convert_texts_to_documents(self.text_split(text_data)) if text_data else []
        self.retriever = None
            
    def convert_texts_to_documents(self, texts):
        return [Document(page_content=text) for text in texts]
        
    def text_split(self , extracted_data):
        self.text_splitter=RecursiveCharacterTextSplitter(chunk_size =200, chunk_overlap =50,separators=["\n\n","\n"])
        return self.text_splitter.split_text(extracted_data)

    def build_retriever(self, connection_string, container_name):
        if not self.documents:
            return
        store = FAISS.from_documents(self.documents, self.embeddings)
        self._save_to_blob(store, connection_string, container_name)
        self.retriever = store.as_retriever()

    def _save_to_blob(self , store, connection_string, container_name):
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        container_client = blob_service_client.get_container_client(container_name)
        with tempfile.NamedTemporaryFile(delete=False) as temp_index:
            faiss.write_index(store.index, temp_index.name)
            with open(temp_index.name, 'rb') as data:
                container_client.upload_blob(name="index.faiss", data=data, overwrite=True)
            os.unlink(temp_index.name)
        with tempfile.NamedTemporaryFile(delete=False) as temp_docstore:
            pickle.dump(store.docstore, temp_docstore)
            temp_docstore.flush()
            with open(temp_docstore.name, 'rb') as data:
                container_client.upload_blob(name="docstore.pkl", data=data, overwrite=True)
            os.unlink(temp_docstore.name)
        with tempfile.NamedTemporaryFile(delete=False) as temp_mapping:
            pickle.dump(store.index_to_docstore_id, temp_mapping)
            temp_mapping.flush()
            with open(temp_mapping.name, 'rb') as data:
                container_client.upload_blob(name="index_to_docstore_id.pkl", data=data, overwrite=True)
            os.unlink(temp_mapping.name)


class RAG_Azure:
    def __init__(self , llm_model="gpt-35-turbo"):
        self.connection_string = os.environ["AZURE_STORAGE_CONNECTION_STRING"]
        self.container_name = os.environ["AZURE_STORAGE_CONTAINER_NAME"]
        self.container_name_historic = os.environ["AZURE_STORAGE_CONTAINER_NAME_HISTORIC"]
        self.client = AzureOpenAI(
          azure_endpoint = os.environ["AZURE_OPENAI_ENDPOINT"],
          api_key= os.environ["AZURE_OPENAI_API_KEY"],
          api_version="2024-05-01-preview"
        )
        self.llm_model=llm_model
        try:
            self.knowledge_base = KnowledgeBase('')
            logger.info(f"knowledge base from empty string is well done")
            self.knowledge_base.retriever = self._load_from_blob().as_retriever()
            logger.info(f"load from blob storage is done. No Building new index.")
        except Exception as e:
            logger.info(f"Could not load from blob storage: {str(e)}. Building new index.")
            self.knowledge_base = KnowledgeBase(self.load_files_contents('data'))
            self.knowledge_base.build_retriever(self.connection_string, self.container_name)
        
    def _load_from_blob(self):
            try:
                logger.info("Démarrage du chargement depuis Azure Blob Storage.")
                logger.info("Connexion à Azure Blob Storage.")
                blob_service_client = BlobServiceClient.from_connection_string(self.connection_string)
                container_client = blob_service_client.get_container_client(self.container_name)
                logger.info("Récupération des blobs 'index.faiss' et 'docstore.pkl'. et index_to_docstore_id.pkl")
                index_blob_client = container_client.get_blob_client("index.faiss")
                docstore_blob_client = container_client.get_blob_client("docstore.pkl")
                docstore_id_blob_client = container_client.get_blob_client("index_to_docstore_id.pkl")
                if not index_blob_client.exists() or not docstore_blob_client.exists() or not docstore_id_blob_client.exists():
                    raise FileNotFoundError("Les blobs nécessaires sont introuvables.")
                index_data = index_blob_client.download_blob().readall()
                docstore_data = docstore_blob_client.download_blob().readall()
                docstore_id_data = docstore_id_blob_client.download_blob().readall()
                logger.info("Chargement de l'index FAISS depuis la mémoire.")
                with tempfile.NamedTemporaryFile(delete=False) as tmp_index:
                        tmp_index.write(index_data)
                        tmp_index_path = tmp_index.name
                index = faiss.read_index(tmp_index_path)
                os.unlink(tmp_index_path)  # Nettoyage immédiat
                logger.info("Chargement du docstore depuis la mémoire.")
                docstore = pickle.loads(docstore_data)
                docstore_id = pickle.loads(docstore_id_data)
                logger.info("Chargement terminé avec succès.")
                return FAISS(self.knowledge_base.embeddings, index, docstore, docstore_id)
            except Exception as e:
                logger.exception(f"Erreur lors du chargement des blobs : {e}")
                raise RuntimeError(f"Erreur lors du chargement des blobs : {e}")
        
    def load_files_contents(self , data_path):
            content_parts = []  
            file_handlers = {
                ".pdf": lambda path: "\n".join(doc.page_content for doc in PyPDFLoader(path).load()),
                ".txt": lambda path: open(path, 'r', encoding='utf-8').read(),
                ".json": lambda path: json.dumps(json.load(open(path, 'r', encoding='utf-8')), indent=2),
                ".docx": lambda path: "\n".join(paragraph.text for paragraph in DocxDocument(path).paragraphs),
            }
            for filename in os.listdir(data_path):
                filepath = os.path.join(data_path, filename)
                file_ext = os.path.splitext(filename)[1]  
                if file_ext in file_handlers:
                    try:
                        content_parts.append(file_handlers[file_ext](filepath))
                    except Exception as e:
                        logger.error(f"Error loading {filepath}: {e}")
            return "\n".join(content_parts)
        
    def answer_query(self, query, context):
        custom_prompt_template = (
                    f"Vous êtes une assistante virtuelle spécialisée exclusivement en imagerie médicale. \n"
                    f"Pour vous aider, voici des informations de contexte:\n {context}\n\n"
                    f"Question du patient : {query} \n\n"
                    f"Si la question sort du domaine de l’imagerie médicale, répondez :« Je suis une assistante spécialisée en imagerie médicale et je ne peux répondre qu'aux questions relatives à ce domaine.»"
                    f"Répondez uniquement à la question posée, utilisez un langage clair, simple et accessible, comme si vous parliez directement à la personne.\n"
                    f"Utilisez le vouvoiement. Soyez concis et précis dans la réponse.\n"       
                )
        try:
                completion = self.client.chat.completions.create(
                    model=self.llm_model,
                    messages=[
                        {"role": "system", "content": custom_prompt_template},
                        {"role": "user", "content": query}
                    ],
                )
                answer = completion.choices[0].message.content
                save_interaction_to_blob(self, question, answer)
                return answer
        except Exception as e:
            logger.error(f"Error answering query: {e}")
            return "Une erreur est survenue lors de la réponse."
            
    def save_interaction_to_blob(self, question, answer):
        try:
            blob_service_client = BlobServiceClient.from_connection_string(self.connection_string)
            container_client = blob_service_client.get_container_client(self.container_name_historic)
            timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S%f")
            blob_name = f"interactions/interaction_{timestamp}.json"
            interaction_data = {
                "question": question,
                "answer": answer,
                "timestamp": timestamp
            }
            json_data = json.dumps(interaction_data, ensure_ascii=False).encode("utf-8")
            container_client.upload_blob(name=blob_name, data=json_data, overwrite=True)
            logger.info(f"Interaction enregistrée sous {blob_name}")
        except Exception as e:
            logger.error(f"Erreur lors de l'enregistrement de l'interaction : {e}")         
            
    def process_query(self, query):
        retrieved_docs = self.knowledge_base.retriever.get_relevant_documents(query , k=2)
        if not retrieved_docs:
            return "Aucun document pertinent trouvé."
        context ="\n".join(doc.page_content for doc in retrieved_docs) 
        return self.answer_query(query, context)

rag_system = RAG_Azure()

def main(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Python HTTP trigger function processed a request.')

    try:
        req_body = req.get_json()
        query = req_body.get('text')

        if not query:
            return func.HttpResponse(
                json.dumps({"error": "No query provided in request body"}),
                mimetype="application/json",
                status_code=400
            )

        result = rag_system.process_query(query)

        return func.HttpResponse(
            json.dumps({"response": result}),
            mimetype="application/json"
        )

    except Exception as e:
        logging.error(f"Error processing request: {str(e)}")
        return func.HttpResponse(
            json.dumps({"error": str(e)}),
            mimetype="application/json",
            status_code=500
        )


